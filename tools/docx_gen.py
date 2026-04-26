import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_horizontal_line(paragraph, color="2E75B6", size=6):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_runs_with_bold(paragraph, text, size=11):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            set_font(run, size=size, bold=True)
        else:
            run = paragraph.add_run(part)
            set_font(run, size=size)


def add_title_page(doc, video_title, video_url):
    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    banner_run = banner.add_run("  ")
    banner_run.font.size = Pt(36)
    pPr = banner._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1F4E79')
    pPr.append(shd)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(video_title)
    set_font(title_run, size=26, bold=True, color="1F4E79")

    doc.add_paragraph()

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_para.add_run("Lecture Study Notes")
    set_font(sub_run, size=16, color="2E75B6")

    doc.add_paragraph()
    doc.add_paragraph()

    divider = doc.add_paragraph()
    add_horizontal_line(divider, color="2E75B6", size=12)

    doc.add_paragraph()

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    set_font(date_run, size=11, color="595959")

    doc.add_paragraph()

    url_para = doc.add_paragraph()
    url_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    url_run = url_para.add_run(f"Source: {video_url}")
    set_font(url_run, size=10, color="2E75B6")

    doc.add_paragraph()
    doc.add_paragraph()

    divider2 = doc.add_paragraph()
    add_horizontal_line(divider2, color="2E75B6", size=12)

    doc.add_page_break()


def create_document(summary: str, video_url: str, video_title: str = "YouTube Lecture") -> dict:
    try:
        doc = Document()

        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)

        add_title_page(doc, video_title, video_url)

        lines = summary.split('\n')

        for line in lines:
            line = line.strip()

            if not line:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                continue

            if line.startswith('## '):
                heading_text = line[3:].strip()
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_after = Pt(4)
                h = doc.add_paragraph()
                h.paragraph_format.space_before = Pt(6)
                h.paragraph_format.space_after = Pt(4)
                run = h.add_run(heading_text.upper())
                set_font(run, size=13, bold=True, color="1F4E79")
                add_horizontal_line(h, color="2E75B6", size=6)

            elif line.startswith('### '):
                heading_text = line[4:].strip()
                h = doc.add_paragraph()
                h.paragraph_format.space_before = Pt(8)
                h.paragraph_format.space_after = Pt(3)
                run = h.add_run(heading_text)
                set_font(run, size=12, bold=True, color="2E75B6")

            elif line.startswith('- ') or line.startswith('* '):
                bullet_text = line[2:].strip()
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.left_indent = Inches(0.3)
                add_runs_with_bold(p, bullet_text, size=11)

            elif len(line) > 2 and line[0].isdigit() and line[1] in '.):':
                p = doc.add_paragraph(style='List Number')
                p.paragraph_format.space_after = Pt(3)
                add_runs_with_bold(p, line[2:].strip(), size=11)

            elif line.startswith('**') and line.endswith('**') and len(line) > 4:
                p = doc.add_paragraph()
                run = p.add_run(line[2:-2])
                set_font(run, size=11, bold=True)

            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                add_runs_with_bold(p, line, size=11)

        for section in doc.sections:
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_horizontal_line(footer_para, color="2E75B6", size=4)
            run = footer_para.add_run(
                f"{video_title}  |  Generated: {datetime.now().strftime('%B %d, %Y')}  |  {video_url}"
            )
            set_font(run, size=8, color="595959")

        output_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "output")
        os.makedirs(output_dir, exist_ok=True)

        safe_title = "".join(c for c in video_title if c.isalnum() or c in (' ', '-', '_')).strip()
        safe_title = safe_title[:50]
        date_str = datetime.now().strftime('%Y%m%d_%H%M')
        filename = f"{safe_title}_Notes_{date_str}.docx"
        filepath = os.path.join(output_dir, filename)

        doc.save(filepath)

        return {
            "success": True,
            "filepath": filepath,
            "filename": filename
        }

    except Exception as e:
        return {
            "success": False,
            "error": f"Document generation failed: {str(e)}"
        }


def add_quiz_to_document(filepath: str, quiz_text: str) -> dict:
    try:
        doc = Document(filepath)

        doc.add_page_break()

        banner = doc.add_paragraph()
        banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr = banner._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F4E79')
        pPr.append(shd)
        run = banner.add_run("  PRACTICE QUIZ  ")
        run.font.size = Pt(20)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string("FFFFFF")

        doc.add_paragraph()

        note = doc.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note_run = note.add_run("Test your understanding — answers are included at the end of each question")
        note_run.font.size = Pt(10)
        note_run.font.color.rgb = RGBColor.from_string("595959")
        note_run.italic = True

        doc.add_paragraph()

        lines = quiz_text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue

            if line.startswith('## '):
                h = doc.add_paragraph()
                h.paragraph_format.space_before = Pt(12)
                run = h.add_run(line[3:].upper())
                run.font.name = "Calibri"
                run.font.size = Pt(13)
                run.bold = True
                run.font.color.rgb = RGBColor.from_string("1F4E79")
                pPr = h._p.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '6')
                bottom.set(qn('w:space'), '1')
                bottom.set(qn('w:color'), '2E75B6')
                pBdr.append(bottom)
                pPr.append(pBdr)

            elif line.startswith('Q') and '.' in line[:4]:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                run = p.add_run(line)
                run.font.name = "Calibri"
                run.font.size = Pt(11)
                run.bold = True

            elif line.startswith('S') and '.' in line[:4]:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                run = p.add_run(line)
                run.font.name = "Calibri"
                run.font.size = Pt(11)
                run.bold = True

            elif line.startswith('Answer:'):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                run = p.add_run(line)
                run.font.name = "Calibri"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor.from_string("2E75B6")
                run.italic = True

            elif len(line) > 2 and line[0] in 'ABCD' and line[1] in ')':
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                run = p.add_run(line)
                run.font.name = "Calibri"
                run.font.size = Pt(11)

            else:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.name = "Calibri"
                run.font.size = Pt(11)

        doc.save(filepath)
        return {"success": True}

    except Exception as e:
        return {"success": False, "error": str(e)}